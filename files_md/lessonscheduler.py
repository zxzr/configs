#!/usr/bin/env python
# -*- coding: utf-8 -*-
"生成授课日期。"
import re
import sys
from docx import Document
from teachings.lesson import Baselesson


def lesson_scheduler():
    """生成授课日期。"""
    try:
        file_md, file_docx = sys.argv[1:3]
    except ValueError:
        print("使用方法：./lessonschedule.py file.md file.docx")
        sys.exit(1)
    print(f"{file_md} >>>>>> {file_docx}")
    with open(file_md, 'r') as t_f:
        new_md = t_f.readlines()
    new_docx = Document(file_docx)
    h1obj = re.compile(r"^#\s+")
    h2obj = re.compile(r"^##\s+")
    h3obj = re.compile(r"^###\s+")
    h4obj = re.compile(r"^####\s+")
    spt = re.compile(r"\s+")
    wk_no = 0
    day_no = 0
    date = ""
    project = ""
    lessons = []
    tbl_no = 0
    row_no = 2
    for tlin in new_md:
        tlin = tlin.strip()
        if re.match(h1obj, tlin):
            if lessons:
                print(f"表{tbl_no}  行{row_no}")
                print(wk_no - 7)
                print(day_no, date)
                print(project)
                print(lessons)
                print("-------------------------------------")
                new_docx.tables[tbl_no].rows[row_no].cells[0].text = date
                new_docx.tables[tbl_no].rows[row_no].cells[1].text = str(
                    wk_no - 7)
                new_docx.tables[tbl_no].rows[row_no].cells[
                    2].text = project + "\n" + "\n".join(lessons)
                row_no = row_no + 1
                if row_no > 7:
                    tbl_no = tbl_no + 1
                    row_no = 2
                lessons = []
            wk_no = int(re.split(spt, tlin, maxsplit=1)[1]) + 7
        elif re.match(h2obj, tlin):
            if lessons:
                print(f"表{tbl_no}  行{row_no}")
                print(wk_no - 7)
                print(day_no, date)
                print(project)
                print(lessons)
                print("-------------------------------------")
                new_docx.tables[tbl_no].rows[row_no].cells[0].text = date
                new_docx.tables[tbl_no].rows[row_no].cells[1].text = str(
                    wk_no - 7)
                new_docx.tables[tbl_no].rows[row_no].cells[
                    2].text = project + "\n" + "\n".join(lessons)
                row_no = row_no + 1
                if row_no > 7:
                    tbl_no = tbl_no + 1
                    row_no = 2
                lessons = []
            day_no = int(re.split(spt, tlin, maxsplit=1)[1])
            date = Baselesson(wk_no, day_no).date.isoformat()
        elif re.match(h3obj, tlin):
            project = re.split(spt, tlin, maxsplit=1)[1]
        elif re.match(h4obj, tlin):
            lessons.append(re.split(spt, tlin, maxsplit=1)[1])

    print(f"表{tbl_no}  行{row_no}")
    print(wk_no - 7)
    print(day_no, date)
    print(project)
    print(lessons)
    print("-------------------------------------")
    new_docx.tables[tbl_no].rows[row_no].cells[0].text = date
    new_docx.tables[tbl_no].rows[row_no].cells[1].text = str(wk_no - 7)
    new_docx.tables[tbl_no].rows[row_no].cells[
        2].text = project + "\n" + "\n".join(lessons)
    new_docx.save('new.docx')


lesson_scheduler()
