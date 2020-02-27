#!/usr/bin/env python
# -*- coding: utf-8 -*-
"生成授课日期。"
import re
from docx import Document
from teachings.lesson import Baselesson


class Single():
    """single lesson."""
    def __init__(self, wk_no=0, day_no=0, date="", project="", contents=None):
        self.wk_no = wk_no
        self.day_no = day_no
        self.date = date
        self.project = project
        self.contents = contents or []


class Lessonscheduler():
    """教学计划。"""
    def __init__(self, file_md, file_docx):
        """init"""
        self.schdls = []
        with open(file_md, 'r') as t_f:
            self.list_md = t_f.readlines()
            self.new_docx = Document(file_docx)

    def solver(self):
        """对象化数据。"""
        h1obj = re.compile(r"^#\s+")
        h2obj = re.compile(r"^##\s+")
        h3obj = re.compile(r"^###\s+")
        h4obj = re.compile(r"^####\s+")
        spt = re.compile(r"\s+")
        sgl = Single()
        for tlin in self.list_md:
            tlin = tlin.strip()
            if re.match(h1obj, tlin):
                self.schdls.append(sgl)
                sgl = Single()
                sgl.wk_no = int(re.split(spt, tlin, maxsplit=1)[1]) + 7
            elif re.match(h2obj, tlin):
                if sgl.contents:
                    self.schdls.append(sgl)
                    sgl = Single(self.schdls[-1].wk_no)
                sgl.day_no = int(re.split(spt, tlin, maxsplit=1)[1])
                sgl.date = Baselesson(sgl.wk_no, sgl.day_no).date.isoformat()
            elif re.match(h3obj, tlin):
                sgl.project = re.split(spt, tlin, maxsplit=1)[1]
            elif re.match(h4obj, tlin):
                sgl.contents.append(re.split(spt, tlin, maxsplit=1)[1])
        self.schdls.append(sgl)
        sgl = Single()
        self.schdls.pop(0)

    def docx(self):
        """fill docx."""
        tbl_no = 0
        row_no = 2
        for tlin in self.schdls:
            self.new_docx.tables[tbl_no].rows[row_no].cells[0].text = tlin.date
            self.new_docx.tables[tbl_no].rows[row_no].cells[1].text = str(
                tlin.wk_no - 7)
            self.new_docx.tables[tbl_no].rows[row_no].cells[
                2].text = tlin.project + "\n" + "\n".join(tlin.contents)
            row_no += 1
            if row_no > 7:
                tbl_no += 1
                row_no = 2
        self.new_docx.save('new.docx')
