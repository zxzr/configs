#!/usr/bin/env python
# -*- coding: utf-8 -*-
#  from docx.enum.text import WD_ALIGN_PARAGRAPH
#  from docx.oxml.ns import qn
#  from docx.shared import Cm, Pt  # Inches
import sys
from isoweek import Week
from docx import Document
if len(sys.argv) < 2:
    print("需要一个docx文件。")
    sys.exit()
newdocx = Document(sys.argv[1])
#  newdocx = Document("2019611工程力学二授课计划.docx")
tbls = newdocx.tables
print("共有%d个表。" % (len(tbls)))
for tb in tbls:
    for i in range(3):
        wk_no = tb.columns[1].cells[2 * i + 2].text
        dt_1 = Week(2020, int(wk_no) + 7).wednesday()  # 1-22 8-29
        dt_2 = Week(2020, int(wk_no) + 7).friday()  # 1-22 8-29
        cell_1 = 2 * i + 2  # 0,1,2==>2,3 4,5 6,7
        cell_2 = 2 * i + 3
        tb.columns[0].cells[cell_1].text = str(dt_1)
        tb.columns[0].cells[cell_2].text = str(dt_2)
        print(wk_no, dt_1)
        print(wk_no, dt_2)
newdocx.save('test.docx')
