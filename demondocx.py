#!/usr/bin/env python
# -*- coding: utf-8 -*-

# for row in table.rows:
# for cell in row.cells:
#     paragraphs = cell.paragraphs
#     for paragraph in paragraphs:
#         for run in paragraph.runs:
#             font = run.font
#             font.size= Pt(30)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt  # Inches

document = Document()

style = document.styles['Normal']
# 设置西文字体
style.font.name = 'Times New Roman'
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 获取段落样式
paragraph_format = style.paragraph_format
# 首行缩进0.74厘米，即2个字符
paragraph_format.first_line_indent = Cm(0.74)

# 设置标题
title = '文档标题'
# title_ = document.add_heading('Document Title', 0)
title_ = document.add_heading(level=0)
# 标题居中
title_.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 添加标题内容
title_run = title_.add_run(title)
# 设置标题字体大小
title_run.font.size = Pt(14)
# 设置标题西文字体
title_run.font.name = 'Times New Roman'
# 设置标题中文字体
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

#  document.add_picture('monty-truth.png', width=Inches(1.25))

records = ((3, '101', 'Spam'), (7, '422', 'Eggs'),
           (4, '631', 'Spam, spam, eggs, and spam'))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')
